import pdfplumber
import docx
import os
from pathlib import Path
from config.settings import UPLOAD_DIR


class ResumeParser:
    """Parses PDF and DOCX files and extracts raw text with metadata."""

    def parse(self, file_path: str) -> dict:
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".pdf":
            return self._parse_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported format: {ext}. Please upload PDF or DOCX.")

    # ── PDF Parser ───────────────────────────────────────────────────────────
    def _parse_pdf(self, file_path: str) -> dict:
        raw_text = ""
        pages = []

        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages.append(text)
                raw_text += text + "\n"

        return {
            "raw_text":    raw_text.strip(),
            "pages":       pages,
            "metadata":    {"total_pages": total_pages},
            "format":      "PDF"
        }

    # ── DOCX Parser ──────────────────────────────────────────────────────────
    def _parse_docx(self, file_path: str) -> dict:
        doc = docx.Document(file_path)
        raw_text = ""
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
                raw_text += para.text + "\n"

        # Extract text from tables too
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        raw_text += cell.text + "\n"

        return {
            "raw_text":    raw_text.strip(),
            "pages":       paragraphs,
            "metadata":    {"total_pages": len(doc.paragraphs)},
            "format":      "DOCX"
        }

    # ── Streamlit File Saver ─────────────────────────────────────────────────
    def save_uploaded_file(self, uploaded_file, save_dir: str = UPLOAD_DIR) -> str:
        os.makedirs(save_dir, exist_ok=True)
        file_path = os.path.join(save_dir, uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        return file_path